"""
tools/docx_generator.py
-------------------------
Word counterpart to pdf_generator.py — same color scheme, same shared
markdown parsing (markdown_ir.py), same cover-page-vs-compact decision, just
built with python-docx.

One real difference from the PDF: python-docx has no layout engine, so it
can't precompute page numbers the way reportlab can. Instead this uses
native Word fields — TOC, PAGE, NUMPAGES — the same mechanism Word itself
uses. Word computes and displays the real values once the document is
opened; if the TOC shows placeholder text instead of entries, the person
just needs to right-click it and choose "Update Field" (Word does this
automatically on print/print-preview in most configurations).
"""

import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .markdown_ir import (
    parse_markdown_blocks, normalize_heading_levels,
    extract_title_and_summary, has_any_heading,
)

# ----------------------------------------------------------------------------
# Color scheme — same values as pdf_generator.py
# ----------------------------------------------------------------------------
NAVY_DARK  = RGBColor(0x13, 0x35, 0x6D)
NAVY       = RGBColor(0x1C, 0x45, 0x87)
NAVY_MID   = RGBColor(0x1C, 0x4E, 0xA0)
TEXT_DARK  = RGBColor(0x22, 0x22, 0x22)
MUTED_GRAY = RGBColor(0x8A, 0x8A, 0x8A)
RULE_GRAY_HEX = "DDDDDD"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

GRADIENT_STOPS = [(78, 131, 216), (28, 78, 160), (19, 53, 109)]


# ----------------------------------------------------------------------------
# Low-level OXML helpers (python-docx has no high-level API for these)
# ----------------------------------------------------------------------------
def _shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_field(paragraph, field_code, placeholder_text=""):
    """Inserts a native Word field (TOC, PAGE, NUMPAGES, ...) into `paragraph`."""
    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = placeholder_text

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)
    return run


def _add_bottom_border(paragraph, color=RULE_GRAY_HEX, size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ----------------------------------------------------------------------------
# Inline markdown -> Word runs (bold/italic/code/citations)
# ----------------------------------------------------------------------------
_INLINE_RE = re.compile(
    r"\*\*(?P<bold>.+?)\*\*"
    r"|(?<!_)_(?P<italic>[^_]+?)_(?!_)"
    r"|`(?P<code>[^`]+)`"
    r"|\[(?P<citation>Document\s*\d+|Source\s*\d+|\d+)\]",
    re.IGNORECASE,
)


def add_inline_runs(paragraph, text):
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("bold") is not None:
            r = paragraph.add_run(m.group("bold"))
            r.bold = True
        elif m.group("italic") is not None:
            r = paragraph.add_run(m.group("italic"))
            r.italic = True
        elif m.group("code") is not None:
            r = paragraph.add_run(m.group("code"))
            r.font.name = "Consolas"
        elif m.group("citation") is not None:
            r = paragraph.add_run(f" [{m.group('citation')}]")
            r.font.size = Pt(7)
            r.font.color.rgb = MUTED_GRAY
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    return paragraph


# ----------------------------------------------------------------------------
# Base style setup
# ----------------------------------------------------------------------------
def _configure_base_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_DARK

    h1 = document.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    h2 = document.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = NAVY_MID
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    h3 = document.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = NAVY_MID
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)


def _gradient_bar(document, width_inches=6.5):
    steps = 24
    stops = GRADIENT_STOPS
    seg = steps // (len(stops) - 1)
    colors = []
    for i in range(len(stops) - 1):
        for s in range(seg):
            t = s / seg
            colors.append(tuple(int(stops[i][k] + (stops[i + 1][k] - stops[i][k]) * t)
                                 for k in range(3)))
    colors.append(stops[-1])

    table = document.add_table(rows=1, cols=len(colors))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_width = Inches(width_inches / len(colors))
    row = table.rows[0]
    row.height = Pt(7)
    for i, cell in enumerate(row.cells):
        cell.width = col_width
        _shade_cell(cell, "%02X%02X%02X" % colors[i])
        cell.paragraphs[0].text = ""
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    return table


def _add_header_footer(document, doc_title):
    """No header/footer on page 1 (the cover/title area); title + rule and
    'Page X of Y' on every later page — same behavior as the PDF."""
    section = document.sections[0]
    section.different_first_page_header_footer = True

    header_p = section.header.paragraphs[0]
    header_p.text = ""
    run = header_p.add_run(doc_title)
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED_GRAY
    _add_bottom_border(header_p)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = footer_p.add_run("Page ")
    prefix.font.size = Pt(8.5)
    prefix.font.color.rgb = MUTED_GRAY
    _add_field(footer_p, "PAGE").font.size = Pt(8.5)
    mid = footer_p.add_run(" of ")
    mid.font.size = Pt(8.5)
    mid.font.color.rgb = MUTED_GRAY
    _add_field(footer_p, "NUMPAGES").font.size = Pt(8.5)

    # First-page header/footer stay blank by default once the flag above is set.


def _add_toc_field(document):
    p = document.add_paragraph()
    _add_field(p, 'TOC \\o "1-3" \\h \\z \\u',
               placeholder_text='Right-click here and choose "Update Field" to build the table of contents.')


# ----------------------------------------------------------------------------
# Block IR -> Word content
# ----------------------------------------------------------------------------
def _render_table(document, rows):
    if not rows:
        return
    ncols = len(rows[0])
    table = document.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx].cells
        for c_idx in range(ncols):
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = cells[c_idx]
            cell.paragraphs[0].text = ""
            run = cell.paragraphs[0].add_run(text)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = WHITE
                _shade_cell(cell, "1C4587")
            else:
                run.font.color.rgb = TEXT_DARK
                if r_idx % 2 == 0:
                    _shade_cell(cell, "F2F5FA")
    document.add_paragraph()  # small spacer after the table


def render_blocks(document, blocks):
    for b in blocks:
        kind = b["type"]
        if kind == "heading":
            h = document.add_heading(level=b["level"])
            add_inline_runs(h, b["text"])
        elif kind == "paragraph":
            p = document.add_paragraph()
            add_inline_runs(p, b["text"])
        elif kind == "bullet_list":
            for item in b["items"]:
                p = document.add_paragraph(style="List Bullet")
                add_inline_runs(p, item)
        elif kind == "number_list":
            for item in b["items"]:
                p = document.add_paragraph(style="List Number")
                add_inline_runs(p, item)
        elif kind == "table":
            _render_table(document, b["rows"])
        elif kind == "hr":
            p = document.add_paragraph()
            _add_bottom_border(p, color=RULE_GRAY_HEX, size=6)


# ----------------------------------------------------------------------------
# Public entry point
# ----------------------------------------------------------------------------
def build_report_docx(
    content,
    path,
    title=None,
    subtitle=None,
    prepared_for=None,
    prepared_by=None,
    date=None,
    executive_summary=None,
):
    """
    Renders `content` (markdown-ish text) into a branded .docx at `path`,
    mirroring build_report_pdf's cover-page-vs-compact behavior.
    """
    title, body, executive_summary = extract_title_and_summary(
        content, title=title, executive_summary=executive_summary
    )
    if date is None:
        date = datetime.now().strftime("%B %d, %Y")

    blocks = normalize_heading_levels(parse_markdown_blocks(body))
    has_headings = has_any_heading(blocks)
    is_full_report = has_headings or any([subtitle, prepared_for, prepared_by, executive_summary])

    document = Document()
    _configure_base_styles(document)
    _add_header_footer(document, title)

    _gradient_bar(document)

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(18)
    title_p.paragraph_format.space_after = Pt(10)
    title_run = title_p.add_run(title)
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(26 if is_full_report else 20)
    title_run.bold = True
    title_run.font.color.rgb = NAVY

    if not is_full_report:
        date_p = document.add_paragraph()
        date_p.paragraph_format.space_after = Pt(14)
        date_run = date_p.add_run(date)
        date_run.font.size = Pt(9.5)
        date_run.font.color.rgb = MUTED_GRAY
        render_blocks(document, blocks)
        document.save(path)
        return path

    # ---- Full cover page ---------------------------------------------------
    if subtitle:
        sub_p = document.add_paragraph()
        sub_p.paragraph_format.space_after = Pt(12)
        sub_run = sub_p.add_run(subtitle)
        sub_run.bold = True
        sub_run.font.size = Pt(12.5)
        sub_run.font.color.rgb = TEXT_DARK

    def _label(text):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = NAVY
        return p

    def _meta_line(text):
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.add_run(text)

    if prepared_for:
        _label("Prepared for:")
        for line in prepared_for.split("\n"):
            _meta_line(line.strip())

    if prepared_by:
        _label("Prepared by:")
        for line in prepared_by.split("\n"):
            _meta_line(line.strip())

    _label("Date:")
    _meta_line(date)

    if executive_summary:
        h = document.add_heading(level=2)
        h.paragraph_format.space_before = Pt(16)
        run = h.add_run("Executive Summary")
        run.font.color.rgb = NAVY
        p = document.add_paragraph()
        add_inline_runs(p, executive_summary)

    toc_heading = document.add_heading(level=2)
    toc_heading.add_run("Table of Contents").font.color.rgb = NAVY
    _add_toc_field(document)

    document.add_page_break()

    # ---- Body ---------------------------------------------------------------
    render_blocks(document, blocks)

    document.save(path)
    return path